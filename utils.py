
import spacy
import re
from collections import Counter
from docx import Document
import json
from io import BytesIO
import matplotlib.pyplot as plt

nlp = spacy.load("en_core_web_sm")

with open("data/skills.json") as f:
    SKILLS = json.load(f)

def extract_resume_text(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
    return text

def extract_keywords_spacy(text):
    doc = nlp(text)
    keywords = set()
    for chunk in doc.noun_chunks:
        cleaned = chunk.text.lower().strip()
        if len(cleaned.split()) >= 1:
            keywords.add(cleaned)
    return list(keywords)

def match_keywords(resume_text, jd_keywords):
    matched = [kw for kw in jd_keywords if kw.lower() in resume_text.lower()]
    missing = [kw for kw in jd_keywords if kw.lower() not in resume_text.lower()]
    return matched, missing

def categorize_skills(keywords):
    hard = [kw for kw in keywords if kw.lower() in SKILLS['hard']]
    soft = [kw for kw in keywords if kw.lower() in SKILLS['soft']]
    return hard, soft

def get_relevant_sentences(text, keywords):
    sentences = text.split("\n")
    relevant = []
    for line in sentences:
        for kw in keywords:
            if kw.lower() in line.lower():
                relevant.append(line)
                break
    return list(set(relevant))

def create_tailored_resume(sentences):
    doc = Document()
    doc.add_heading("Tailored Resume", 0)
    for s in sentences:
        doc.add_paragraph(s, style="List Bullet")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def draw_pie_chart(matched, missing):
    fig, ax = plt.subplots()
    labels = 'Matched', 'Missing'
    sizes = [matched, missing]
    colors = ['green', 'red']
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    ax.axis('equal')
    return fig

def calculate_resume_score(matched_keywords, total_keywords, relevant_sentences, total_sentences, hard_skills, soft_skills):
    if total_keywords == 0: return 0
    skill_match_ratio = len(matched_keywords) / total_keywords
    context_ratio = len(relevant_sentences) / (total_sentences if total_sentences > 0 else 1)
    skill_balance = min(len(hard_skills), len(soft_skills)) / max(len(hard_skills) + len(soft_skills), 1)
    score = (
        (skill_match_ratio * 0.5) + 
        (context_ratio * 0.2) +
        (min(len(matched_keywords) / 15, 1.0) * 0.2) +
        (skill_balance * 0.1)
    ) * 100
    return round(score, 2)
